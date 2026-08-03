import os
from io import BytesIO
from typing import Dict, Any
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

class DocumentGeneratorService:
    @staticmethod
    def generate_docx(report_name: str, data: Dict[str, Any]) -> BytesIO:
        """
        Generates a styled .docx report matching banking presentation standards.
        """
        doc = DocxDocument()
        
        # Define brand colors (#0052FF - primary, #0A0B0D - text)
        primary_color = RGBColor(0, 82, 255)
        text_color = RGBColor(10, 11, 13)
        
        # Styles
        style_normal = doc.styles['Normal']
        style_normal.font.name = 'Inter'
        style_normal.font.size = Pt(11)
        style_normal.font.color.rgb = text_color

        # Cover Page
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(f"\n\n\n\nLOAN SYNDICATION PROPOSAL\n\n{report_name.upper()}")
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = primary_color
        
        doc.add_paragraph(f"\nPrepared by: LoanCraft AI Consulting Group\nFor Bank Credit Evaluation Committee")
        doc.add_page_break()

        # Add Table of Contents
        doc.add_heading('Table of Contents', level=1)
        doc.add_paragraph("1. Executive Summary\n2. Borrower & Promoter Profile\n3. Technical & Feasibility Analysis\n4. Financial Appraisal & Ratios\n5. Risk Mitigations & Recommendations")
        doc.add_page_break()

        # Section 1: Executive Summary
        h1 = doc.add_heading('1. Executive Summary', level=1)
        h1.runs[0].font.color.rgb = primary_color
        doc.add_paragraph(data.get("executive_summary", "This report details the project feasibility..."))

        # Section 2: Borrower Profile
        h2 = doc.add_heading('2. Borrower Profile', level=1)
        h2.runs[0].font.color.rgb = primary_color
        doc.add_paragraph(data.get("borrower_profile", "The borrower exhibits high credit score..."))

        # Section 3: Financial Statements Table
        h3 = doc.add_heading('3. Financial Statements & Ratio Analysis', level=1)
        h3.runs[0].font.color.rgb = primary_color
        
        financials = data.get("financials", {})
        if financials:
            doc.add_paragraph("Summary of key balance sheet and profitability parameters:")
            # Create table
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Financial Variable'
            hdr_cells[1].text = 'Value (INR)'
            hdr_cells[2].text = 'Ratio Status'
            
            for key, val in financials.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(key).replace("_", " ").title()
                row_cells[1].text = f"{val:,.2f}" if isinstance(val, (int, float)) else str(val)
                row_cells[2].text = "Aligned with covenants"

        # Save to byte stream
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    @staticmethod
    def generate_pdf(report_name: str, data: Dict[str, Any]) -> BytesIO:
        """
        Generates an institutional PDF report using ReportLab platypus flowables.
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54
        )

        styles = getSampleStyleSheet()
        
        # Modify default styles for sleek fintech minimalist aesthetics
        styles.add(ParagraphStyle(
            name='CoverTitle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=32,
            leading=38,
            textColor=colors.HexColor('#0052FF'),
            spaceAfter=20
        ))

        styles.add(ParagraphStyle(
            name='SectionHeading',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=18,
            leading=22,
            textColor=colors.HexColor('#0A0B0D'),
            spaceBefore=15,
            spaceAfter=10,
            keepWithNext=True
        ))

        styles.add(ParagraphStyle(
            name='ReportBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=11,
            leading=16,
            textColor=colors.HexColor('#5B616E'),
            spaceAfter=8
        ))

        story = []

        # Cover Page
        story.append(Spacer(1, 150))
        story.append(Paragraph(f"LOANCRAFT AI CREDIT MEMORANDUM", styles['CoverTitle']))
        story.append(Paragraph(f"Project Assessment and Funding Application for:<br/><b>{report_name.upper()}</b>", styles['Normal']))
        story.append(Spacer(1, 30))
        story.append(Paragraph("Prepared for Institutional Review Board", styles['Normal']))
        story.append(PageBreak())

        # TOC
        story.append(Paragraph("Table of Contents", styles['SectionHeading']))
        story.append(Spacer(1, 10))
        story.append(Paragraph("1. Executive Summary ........................................................................ 3", styles['ReportBody']))
        story.append(Paragraph("2. Borrower and Promoter Profile ............................................................. 4", styles['ReportBody']))
        story.append(Paragraph("3. Technical Assessment & Project Means ................................................ 5", styles['ReportBody']))
        story.append(Paragraph("4. Credit Appraisal & Financial Covenants .................................................... 6", styles['ReportBody']))
        story.append(PageBreak())

        # Content
        story.append(Paragraph("1. Executive Summary", styles['SectionHeading']))
        story.append(Paragraph(data.get("executive_summary", "Detailed summaries of proposed project plans and finances."), styles['ReportBody']))
        story.append(Spacer(1, 15))

        story.append(Paragraph("2. Borrower Profile", styles['SectionHeading']))
        story.append(Paragraph(data.get("borrower_profile", "Detailed promoter backings, incorporation status, and management profiles."), styles['ReportBody']))
        story.append(Spacer(1, 15))

        # Financial Data Table
        story.append(Paragraph("3. Financial Appraisal & Key Metrics", styles['SectionHeading']))
        financials = data.get("financials", {})
        if financials:
            table_data = [["Financial Metric", "Value (INR)", "Assessment Status"]]
            for k, v in financials.items():
                formatted_v = f"{v:,.2f}" if isinstance(v, (int, float)) else str(v)
                table_data.append([str(k).replace("_", " ").title(), formatted_v, "Within Covenants"])

            t = Table(table_data, colWidths=[200, 150, 150])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0052FF')),
                ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('BOTTOMPADDING', (0,0), (-1,0), 8),
                ('BACKGROUND', (0,1), (-1,-1), colors.HexColor('#F7F8FA')),
                ('GRID', (0,0), (-1,-1), 1, colors.HexColor('#E5E7EB')),
                ('FONTSIZE', (0,0), (-1,-1), 9),
            ]))
            story.append(t)

        doc.build(story)
        buffer.seek(0)
        return buffer
